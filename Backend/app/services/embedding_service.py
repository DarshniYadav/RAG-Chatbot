from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from app.utils.vector_store import get_vector_store
import os
from pathlib import Path
from datetime import datetime
from docx import Document as DocxDocument


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}


def _load_documents(file_path: str):
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return PyPDFLoader(file_path).load()

    if ext in {".txt", ".md"}:
        return TextLoader(file_path, encoding="utf-8").load()

    if ext == ".docx":
        docx = DocxDocument(file_path)
        content = "\n".join([p.text for p in docx.paragraphs if p.text and p.text.strip()])
        if not content.strip():
            content = ""
        return [Document(page_content=content, metadata={"source": file_path})]

    raise ValueError(f"Unsupported file type '{ext}'. Allowed: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")

async def process_document(file_path: str, user_id: str | None = None, original_filename: str | None = None):
    try:
        docs = _load_documents(file_path)

        # Use smaller chunks to better separate different sections/people in documents
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,  # Reduced from 1000
            chunk_overlap=150,  # Reduced from 200
            separators=["\n\n", "\n", " ", ""]  # Prefer paragraph breaks
        )
        splits = splitter.split_documents(docs)

        # Add document metadata to each chunk
        filename = Path(file_path).name
        display_name = original_filename or filename
        timestamp = datetime.utcnow().isoformat()
        
        for i, doc in enumerate(splits):
            # Preserve existing metadata and add source info
            if not doc.metadata:
                doc.metadata = {}
            doc.metadata["source_file"] = filename
            doc.metadata["source_display_name"] = display_name
            doc.metadata["chunk_index"] = i
            doc.metadata["total_chunks"] = len(splits)
            doc.metadata["upload_time"] = timestamp
            doc.metadata["page"] = doc.metadata.get("page", 0)
            if user_id:
                doc.metadata["user_id"] = str(user_id)

        vectorstore = get_vector_store()
        vectorstore.add_documents(splits)

    finally:
        # Clean up uploaded temp file whether processing succeeds or fails.
        if os.path.exists(file_path):
            os.remove(file_path)